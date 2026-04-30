# parsers/document_parser.py
"""
Document Parser - parse PDF, DOCX, CSV, Excel, TXT into LangChain Documents.
FIX: uses pypdf instead of deprecated PyPDF2.
"""

import pandas as pd
from pathlib import Path
from typing import List

from langchain.schema import Document

# FIX: pypdf replaces deprecated PyPDF2 (better extraction, no known bugs)
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader  # graceful fallback if pypdf not installed

from docx import Document as DocxDocument
from utils.logger import logger


class DocumentParser:
    """Parse multiple document formats into LangChain Documents."""

    @staticmethod
    def parse_pdf(file_path: str) -> List[Document]:
        documents = []
        filename  = Path(file_path).name
        logger.info(f"📕 Parsing PDF: {filename}")

        reader      = PdfReader(file_path)
        total_pages = len(reader.pages)

        documents.append(Document(
            page_content=f"PDF Document: {filename}\nTotal Pages: {total_pages}",
            metadata={"source": filename, "type": "pdf_summary", "total_pages": total_pages},
        ))

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(Document(
                    page_content=f"From {filename}, Page {page_num + 1}:\n{text.strip()}",
                    metadata={
                        "source": filename, "type": "pdf",
                        "page": page_num + 1, "total_pages": total_pages,
                    },
                ))

        logger.info(f"✅ Parsed {len(documents)} chunks from {filename}")
        return documents

    @staticmethod
    def parse_docx(file_path: str) -> List[Document]:
        doc      = DocxDocument(file_path)
        documents = []
        filename  = Path(file_path).name
        logger.info(f"📝 Parsing Word Doc: {filename}")

        para_count  = sum(1 for p in doc.paragraphs if p.text.strip())
        table_count = len(doc.tables)

        documents.append(Document(
            page_content=(
                f"Word Document: {filename}\n"
                f"Paragraphs: {para_count}\nTables: {table_count}"
            ),
            metadata={
                "source": filename, "type": "docx_summary",
                "paragraph_count": para_count, "table_count": table_count,
            },
        ))

        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                documents.append(Document(
                    page_content=f"From {filename}, Paragraph {idx}:\n{para.text.strip()}",
                    metadata={"source": filename, "type": "docx_paragraph", "paragraph": idx},
                ))

        for tidx, table in enumerate(doc.tables):
            rows = "\n".join(
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            )
            documents.append(Document(
                page_content=f"From {filename}, Table {tidx}:\n{rows}",
                metadata={"source": filename, "type": "docx_table", "table": tidx},
            ))

        logger.info(f"✅ Parsed {len(documents)} chunks from {filename}")
        return documents

    @staticmethod
    def parse_txt(file_path: str) -> List[Document]:
        filename = Path(file_path).name
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        return [Document(
            page_content=text,
            metadata={"source": filename, "type": "txt"},
        )]

    @staticmethod
    def parse_csv(file_path: str) -> List[Document]:
        documents = []
        df        = pd.read_csv(file_path)
        filename  = Path(file_path).name
        logger.info(f"📄 Parsing CSV: {filename}")

        summary = (
            f"CSV File: {filename}\n"
            f"Columns: {', '.join(df.columns)}\n"
            f"Total Rows: {len(df)}\n\n"
            f"Sample Data:\n{df.head(5).to_string()}"
        )
        documents.append(Document(
            page_content=summary,
            metadata={"source": filename, "type": "csv_summary",
                      "row_count": len(df), "columns": list(df.columns)},
        ))

        for idx, row in df.iterrows():
            content = f"From {filename}, Row {idx}:\n"
            content += ", ".join(f"{c}: {v}" for c, v in row.items())
            documents.append(Document(
                page_content=content,
                metadata={"source": filename, "type": "csv_row", "row": int(idx)},
            ))

        logger.info(f"✅ Parsed {len(documents)} chunks from {filename}")
        return documents

    @staticmethod
    def parse_excel(file_path: str) -> List[Document]:
        documents  = []
        excel_file = pd.ExcelFile(file_path)
        filename   = Path(file_path).name
        logger.info(f"📊 Parsing Excel: {filename}")

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            summary = (
                f"Excel File: {filename}, Sheet: {sheet_name}\n"
                f"Columns: {', '.join(df.columns)}\nTotal Rows: {len(df)}\n\n"
                f"Sample:\n{df.head(3).to_string()}"
            )
            documents.append(Document(
                page_content=summary,
                metadata={"source": filename, "type": "excel",
                          "sheet": sheet_name, "row_count": len(df)},
            ))

            for idx, row in df.iterrows():
                content = f"From {filename}, Sheet '{sheet_name}', Row {idx}:\n"
                content += ", ".join(f"{c}: {v}" for c, v in row.items())
                documents.append(Document(
                    page_content=content,
                    metadata={"source": filename, "type": "excel_row",
                              "sheet": sheet_name, "row": int(idx)},
                ))

        logger.info(f"✅ Parsed {len(documents)} chunks from {filename}")
        return documents